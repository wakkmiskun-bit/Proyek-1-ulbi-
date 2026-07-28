import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx library not found. Please install it with 'pip install python-docx'.")
    sys.exit(1)

def build_docx():
    doc = Document()
    
    # Page setup (Standard A4 Margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # Helper to add heading
    def add_heading(text, level, color=RGBColor(30, 58, 138)): # Royal Blue
        p = doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.runs[0]
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.color.rgb = color
        return p
        
    # Helper to add bullet
    def add_bullet(text, bold_prefix="", level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.bold = True
            run_bold.font.name = 'Calibri'
            run_bold.font.size = Pt(11)
            
        run_text = p.add_run(text)
        run_text.font.name = 'Calibri'
        run_text.font.size = Pt(11)
        return p
        
    # Helper to add presenter speech block
    def add_speech_block(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        
        # Style speech text differently (italic, slate gray)
        run = p.add_run(f'Presenter Script:\n"{text}"')
        run.italic = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        return p

    # Cover / Header Details
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NASKAH PRESENTASI SIDANG PROYEK 1\nTASKMATE - KELOMPOK 22\n")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Sistem Manajemen Tugas dan Produktivitas Berbasis Web\nUniversitas Logistik & Bisnis Internasional (ULBI)")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    p_sub.paragraph_format.space_after = Pt(24)
    
    # Meta Info Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    meta_info = [
        ("Nama Kelompok", "Kelompok 22"),
        ("Anggota Kelompok", "1. Muhammad Ilham Habiballah (NPM: 714250003)\n2. Gianjar Nugraha (NPM: 714250007)"),
        ("Program Studi", "D4 Teknik Informatika - ULBI"),
        ("Dosen Pembimbing", "Cahyo Prianto, S.Pd., M.T., CDSP, SFPC"),
        ("Tahun Akademik", "2025/2026")
    ]
    for i, (k, v) in enumerate(meta_info):
        row = table.rows[i]
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Slide 1
    add_heading("SLIDE 1: COVER (PEMBUKAAN)", level=1)
    add_bullet(" Judul Aplikasi: TaskMate - Sistem Layanan Manajemen Tugas & Produktivitas Berbasis Web.", bold_prefix="Fokus:")
    add_bullet(" Memperkenalkan seluruh tim Kelompok 22 dan Dosen Pembimbing Bapak Cahyo Prianto.", bold_prefix="Tujuan:")
    add_speech_block(
        "Assalamu’alaikum Warahmatullahi Wabarakatuh. Selamat pagi/siang kepada Bapak/Ibu Dewan Penguji dan Dosen Pembimbing Bapak Cahyo Prianto. "
        "Perkenalkan, kami dari Kelompok 22 Program Studi D4 Teknik Informatika, Universitas Logistik dan Bisnis Internasional. "
        "Pada kesempatan hari ini, kami akan mempresentasikan hasil pengembangan aplikasi Proyek 1 kami yang berjudul 'TaskMate: Sistem Manajemen Tugas dan Produktivitas Berbasis Web'. "
        "Saya Muhammad Ilham Habiballah dan rekan saya Gianjar Nugraha akan menjelaskan bagaimana aplikasi ini dapat memecahkan masalah produktivitas akademik di lingkungan mahasiswa ULBI. Mari kita mulai ke pembahasan latar belakang."
    )

    # Slide 2
    add_heading("SLIDE 2: LATAR BELAKANG & TUJUAN", level=1)
    add_bullet(" Padatnya kurikulum praktikum ULBI, kelalaian deadline tugas (forgetfulness), ketidaktersediaan monitoring dari admin akademik.", bold_prefix="Masalah:")
    add_bullet(" Membangun platform Kanban visual terintegrasi WhatsApp Gateway Fonnte API untuk notifikasi proaktif.", bold_prefix="Tujuan:")
    add_speech_block(
        "Bapak dan Ibu Penguji, pengembangan TaskMate dilatarbelakangi oleh permasalahan kesibukan akademik mahasiswa D4 TI ULBI. "
        "Pencatatan tugas harian konvensional bersifat pasif dan rentan terlupakan. Oleh karena itu, tujuan solutif dari TaskMate is "
        "menyediakan platform manajemen tugas terpusat berbasis web. Di sini kami merancang 4 pilar fitur utama: Papan Kanban Visual untuk "
        "pengelolaan dinamis, WhatsApp Gateway untuk pengingat proaktif (H-5 & H-2 sebelum deadline), Multi-Guard Auth, dan Dasbor Monitoring."
    )

    # Slide 3
    add_heading("SLIDE 3: FITUR UNGGULAN APLIKASI", level=1)
    add_bullet(" Pemisahan login Mahasiswa (web) & Admin (admin guard).", bold_prefix="Multi-Role Auth:")
    add_bullet(" Kolom To Do, Doing, Review, Done dengan SortableJS (drag-drop) & auto-done 100%.", bold_prefix="Kanban Board:")
    add_bullet(" Pengiriman pengingat tenggat otomatis ke WA via Fonnte API (scheduler cron).", bold_prefix="WhatsApp Gateway:")
    add_bullet(" Grafik keaktifan mahasiswa dan log audit trail aktivitas sistem.", bold_prefix="Monitoring & Audit:")
    add_speech_block(
        "Aplikasi TaskMate memiliki 4 fitur unggulan utama. Pertama, Multi-Role Authorization untuk menjaga isolasi hak akses. "
        "Kedua, Visual Kanban Board dengan interaksi drag-and-drop mulus serta status auto-done. Ketiga, WhatsApp Gateway Reminder "
        "menggunakan scheduler Fonnte API untuk mengirim pengingat otomatis pada H-5 dan H-2. Keempat, Live Monitoring & Audit Log "
        "di dasbor Admin untuk transparansi aktivitas sistem."
    )

    # Slide 4
    add_heading("SLIDE 4: ALUR KERJA OPERASIONAL SISTEM (WORKFLOW)", level=1)
    add_bullet(" Input Tugas -> Drag & Drop Kanban -> Checklist Sub-tugas 100% -> Done & Terima WhatsApp.", bold_prefix="Mahasiswa:")
    add_bullet(" Pantau Dasbor Global -> Live Board Monitoring -> Audit Log -> Trigger Command WhatsApp.", bold_prefix="Admin:")
    add_speech_block(
        "Alur kerja kami berjalan sangat efisien. Mahasiswa menginput tugas beserta checklist sub-tugasnya. "
        "Mahasiswa menggeser tugas sesuai progresnya. Ketika checklist 100%, status otomatis bergeser ke DONE. "
        "Dari sisi Admin, Admin memantau dasbor global dan live board mahasiswa secara langsung, serta memeriksa audit log aktivitas."
    )

    # Slide 5
    add_heading("SLIDE 5: KESIMPULAN & INTEGRASI WHATSAPP", level=1)
    add_bullet(" Digitalisasi papan kerja Kanban terbukti mempermudah prioritas tugas akademik.", bold_prefix="Kesimpulan 1:")
    add_bullet(" Notifikasi WhatsApp proaktif sukses meminimalkan keterlambatan pengumpulan.", bold_prefix="Kesimpulan 2:")
    add_bullet(" Dasbor multi-role dan Help Center mempercepat aduan teknis ke Admin.", bold_prefix="Kesimpulan 3:")
    add_speech_block(
        "Sebagai penutup, aplikasi TaskMate v2.0 sukses dirancang dan diimplementasikan 100%. Di sebelah kanan slide "
        "dapat dilihat bukti nyata chat pengingat WhatsApp yang dikirim otomatis oleh scheduler server menggunakan Fonnte API. "
        "Sistem ini siap diimplementasikan untuk meningkatkan kedisiplinan belajar mahasiswa ULBI. Terima kasih atas perhatian Dewan Penguji."
    )

    # Q&A Section
    add_heading("PREDIKSI PERTANYAAN & JAWABAN (Q&A)", level=1, color=RGBColor(220, 38, 38)) # Red Accent
    
    qa_list = [
        ("Mengapa memilih WhatsApp Gateway dibanding email?", 
         "WhatsApp memiliki tingkat keterbacaan (open-rate) hampir 98% dan dibaca secara real-time oleh mahasiswa dibandingkan email yang jarang dibuka."),
        ("Bagaimana scheduler notifikasi bekerja?", 
         "Menggunakan Cron Job di server yang mengeksekusi perintah 'php artisan schedule:run' secara berkala. Perintah ini mengecek data tenggat waktu di database lalu mengirimkan pesan lewat API cURL Fonnte ke WhatsApp tujuan."),
        ("Bagaimana keamanan data antar mahasiswa terjamin?", 
         "Setiap mahasiswa login melalui Laravel web guard terpisah. Query database diisolasi berdasarkan 'mahasiswa_id' aktif dalam session, sehingga tidak ada kebocoran data antar mahasiswa.")
    ]
    
    for q, a in qa_list:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"❓ Pertanyaan: {q}")
        run_q.bold = True
        run_q.font.color.rgb = RGBColor(220, 38, 38)
        p_q.paragraph_format.space_before = Pt(8)
        
        p_a = doc.add_paragraph()
        p_a.add_run("💡 Jawaban: ").bold = True
        p_a.add_run(a)
        p_a.paragraph_format.space_after = Pt(8)

    output_file = "Naskah_Presentasi_Kelompok_22.docx"
    doc.save(output_file)
    print(f"Word document saved successfully to: {os.path.abspath(output_file)}")

if __name__ == "__main__":
    build_docx()
