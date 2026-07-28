import os
import sys

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx not found.")
    sys.exit(1)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_naskah_word():
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    COLOR_PRIMARY = RGBColor(233, 30, 99)   # Pink (#e91e63)
    COLOR_DARK = RGBColor(26, 26, 46)      # Navy (#1a1a2e)
    COLOR_ILHAM = RGBColor(37, 99, 235)    # Blue (#2563eb)
    COLOR_GIANJAR = RGBColor(142, 36, 170) # Purple (#8e24aa)
    COLOR_TEXT = RGBColor(40, 40, 40)
    COLOR_MUTED = RGBColor(100, 100, 100)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT

    # Header Title Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("NASKAH PRESENTASI & HAFALAN SIDANG")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("TaskMate: Sistem Manajemen Tugas dan Produktivitas Berbasis Web\nKelompok 26 — D4 Teknik Informatika ULBI 2026")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(12)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Metadata Table
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Nama Aplikasi", "TaskMate (Sistem Manajemen Tugas Berbasis Web)"),
        ("Anggota Kelompok 26", "1. Muhammad Ilham Habiballah (NPM: 714250003)\n2. Gianjar Nugraha (NPM: 714250007)"),
        ("Dosen Pembimbing", "Cahyo Prianto, S.Pd., M.T., CDSP, SFPC (NIK: 117.84.222)"),
        ("Institusi & Akademik", "D4 Teknik Informatika - ULBI (Tahun Akademik 2025/2026)")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        cell_l, cell_r = row.cells[0], row.cells[1]
        cell_l.width, cell_r.width = Inches(2.2), Inches(4.3)
        set_cell_background(cell_l, "F5F7FA")
        set_cell_margins(cell_l, 100, 100, 150, 150)
        set_cell_margins(cell_r, 100, 100, 150, 150)

        p_l = cell_l.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.font.bold = True
        r_l.font.size = Pt(10)
        r_l.font.color.rgb = COLOR_DARK

        p_r = cell_r.paragraphs[0]
        r_r = p_r.add_run(val)
        r_r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # RANGKUMAN PEMBAGIAN SLIDE
    h_sum = doc.add_heading(level=1)
    r_h = h_sum.add_run("📌 RANGKUMAN PEMBAGIAN HAFALAN SLIDE (ILHAM vs GIANJAR)")
    r_h.font.name = "Arial"
    r_h.font.size = Pt(14)
    r_h.font.bold = True
    r_h.font.color.rgb = COLOR_PRIMARY

    p_guide = doc.add_paragraph()
    r_g = p_guide.add_run("Struktur urutan slide disesuaikan secara tepat dengan alur yang Anda minta (Latar Belakang -> Tujuan -> Masalah -> Tampilan -> Database -> Alur Kerja & Bantuan -> Hasil & Simpulan):")
    r_g.font.italic = True
    r_g.font.size = Pt(10.5)

    table_split = doc.add_table(rows=10, cols=3)
    table_split.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table_split.rows[0].cells
    hdr_cells[0].text = "Slide #"
    hdr_cells[1].text = "Topik Utama PPT (Dengan Footer)"
    hdr_cells[2].text = "Pembicara (Siapa yang Menjelaskan)"

    for c in hdr_cells:
        set_cell_background(c, "E91E63")
        set_cell_margins(c, 120, 120, 150, 150)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    slides_summary = [
        ("Slide 1", "Cover & Pembukaan Sidang", "MUHAMMAD ILHAM HABIBALLAH"),
        ("Slide 2", "1. Latar Belakang Pengembangan", "MUHAMMAD ILHAM HABIBALLAH"),
        ("Slide 3", "2. Tujuan Pengembangan (Umum & Khusus)", "GIANJAR NUGRAHA"),
        ("Slide 4", "3. Masalah Akademik Utama", "GIANJAR NUGRAHA"),
        ("Slide 5", "4. Tampilan Aplikasi & Modul UI", "MUHAMMAD ILHAM HABIBALLAH"),
        ("Slide 6", "5. Database MySQL / Schema ERD (Sunnah)", "MUHAMMAD ILHAM HABIBALLAH"),
        ("Slide 7", "6. Alur Kerja Aplikasi & Layanan Pusat Bantuan", "GIANJAR NUGRAHA (SLIDE UTAMA)"),
        ("Slide 8", "7. Hasil Akhir & Kesimpulan", "MUHAMMAD ILHAM HABIBALLAH"),
        ("Slide 9", "Penutup, Q&A & Demo Live Aplikasi", "ILHAM & GIANJAR (BERSAMA-SAMA)")
    ]

    for idx, (snum, stopic, sspeaker) in enumerate(slides_summary):
        row = table_split.rows[idx + 1]
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]
        c0.width, c1.width, c2.width = Inches(1.0), Inches(3.2), Inches(2.3)
        c0.text, c1.text, c2.text = snum, stopic, sspeaker

        bg = "F9FAFB" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_background(c2, bg)

        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)
        set_cell_margins(c2, 80, 80, 120, 120)

        p2 = c2.paragraphs[0]
        for r in p2.runs:
            r.font.bold = True
            if "ILHAM" in sspeaker:
                r.font.color.rgb = COLOR_ILHAM
            elif "GIANJAR" in sspeaker:
                r.font.color.rgb = COLOR_GIANJAR

    doc.add_page_break()

    # DETAIL NASKAH PER SLIDE
    h_detail = doc.add_heading(level=1)
    r_hd = h_detail.add_run("🗣️ DETAIL NASKAH DIALOG PER SLIDE (TEKS HAFALAN)")
    r_hd.font.name = "Arial"
    r_hd.font.size = Pt(16)
    r_hd.font.bold = True
    r_hd.font.color.rgb = COLOR_PRIMARY

    slides_detail = [
        {
            "num": "SLIDE 1: COVER & PEMBUKAAN SIDANG",
            "speaker": "MUHAMMAD ILHAM HABIBALLAH",
            "speaker_color": COLOR_ILHAM,
            "script": (
                "\"Assalamu’alaikum Warahmatullahi Wabarakatuh. Selamat pagi/siang kami ucapkan kepada "
                "Bapak/Ibu Dewan Penguji dan Dosen Pembimbing kami, Bapak Cahyo Prianto.\n\n"
                "Perkenalkan, kami dari Kelompok 26. Saya Muhammad Ilham Habiballah dan rekan saya Gianjar Nugraha. "
                "Pada hari ini, kami akan mempresentasikan hasil pengembangan Proyek 1 kami yang berjudul "
                "'TaskMate: Sistem Manajemen Tugas dan Produktivitas Berbasis Web'.\n\n"
                "Mari kita mulai dari poin pertama yaitu Latar Belakang Pengembangan.\""
            ),
            "tip": "Ucapkan salam dengan sopan, berdiri tegak, dan tatap Dewan Penguji dengan percaya diri."
        },
        {
            "num": "SLIDE 2: 1. LATAR BELAKANG PENGEMBANGAN",
            "speaker": "MUHAMMAD ILHAM HABIBALLAH",
            "speaker_color": COLOR_ILHAM,
            "script": (
                "\"Bapak/Ibu Penguji yang kami hormati, sebagai mahasiswa D4 Teknik Informatika ULBI, kami menghadapi "
                "jadwal perkuliahan yang padat, praktikum laboratorium, dan proyek kelompok yang tenggat waktunya sering berdekatan.\n\n"
                "Sayangnya, sebagian besar mahasiswa masih mengandalkan pencatatan manual di buku atau catatan HP seadanya "
                "yang terpisah dan tidak memiliki pengingat aktif.\n\n"
                "Di sisi lain, pihak pengelola akademik atau Admin juga kesulitan memantau keaktifan mahasiswa secara real-time.\n\n"
                "Peluang teknologi modern seperti Laravel dan Tailwind CSS mendorong kami membangun aplikasi TaskMate yang "
                "memadukan papan kerja visual Kanban dan pengingat notifikasi otomatis via WhatsApp Gateway.\""
            ),
            "tip": "Jelaskan konteks kesibukan mahasiswa ULBI dan pentingnya kehadiran sistem terintegrasi."
        },
        {
            "num": "SLIDE 3: 2. TUJUAN PENGEMBANGAN (UMUM & KHUSUS)",
            "speaker": "GIANJAR NUGRAHA",
            "speaker_color": COLOR_GIANJAR,
            "script": (
                "\"Melanjutkan penjelasan rekan saya Ilham, Tujuan Umum dari proyek ini adalah membangun aplikasi TaskMate berbasis web "
                "menggunakan Laravel, Tailwind CSS, dan MySQL untuk meningkatkan efisiensi waktu dan produktivitas mahasiswa ULBI.\n\n"
                "Adapun Tujuan Khusus yang telah kami selesaikan 100% meliputi:\n"
                "• Hak akses Multi-role terisolasi antara Admin dan Mahasiswa.\n"
                "• Fitur CRUD Task lengkap dengan prioritas dan tanggal deadline.\n"
                "• Modul Kanban Board visual 4 kolom (To Do, Doing, Review, Done).\n"
                "• Integrasi WhatsApp Service API Fonnte untuk pengiriman naskah pengingat H-5 dan H-2 secara manual.\n"
                "• Dasbor Admin Monitoring dengan statistik real-time dan audit log aktivitas.\""
            ),
            "tip": "Sebutkan 5 luaran khusus dan tegaskan bahwa seluruh fitur telah rampung 100%."
        },
        {
            "num": "SLIDE 4: 3. PERMASALAHAN UTAMA",
            "speaker": "GIANJAR NUGRAHA",
            "speaker_color": COLOR_GIANJAR,
            "script": (
                "\"Bapak/Ibu Penguji, permasalahan utama yang kami identifikasi meliputi 4 poin penting:\n\n"
                "1. Kelupaan dan Menunda Pekerjaan (Prokrastinasi): Banyaknya deadline membuat mahasiswa sering lupa dan tugas menumpuk di akhir semester.\n"
                "2. Tidak Ada Pengingat Otomatis: Catatan biasa tidak memiliki sistem notifikasi proaktif sebelum batas waktu habis.\n"
                "3. Kesulitan Menentukan Prioritas: Tanpa visualisasi Kanban, mahasiswa bingung memilah tugas yang baru masuk, sedang dikerjakan, atau sudah selesai.\n"
                "4. Minim Monitoring Admin: Admin tidak memiliki visibilitas dan audit log untuk memantau produktivitas tugas mahasiswa.\""
            ),
            "tip": "Tekankan 4 masalah utama akademik dengan tegas agar penguji memahami urgency solusi TaskMate."
        },
        {
            "num": "SLIDE 5: 4. TAMPILAN & MODUL ANTARMUKA APLIKASI",
            "speaker": "MUHAMMAD ILHAM HABIBALLAH",
            "speaker_color": COLOR_ILHAM,
            "script": (
                "\"Dari sisi antarmuka, TaskMate mengusung konsep Clean & Glassmorphism dengan modul utama sebagai berikut:\n\n"
                "Di sisi Mahasiswa, terdapat halaman Dashboard Workspace (rekapitulasi data, progress bar, tugas terdekat deadline, kalender), Papan Kanban Interaktif 4 status dengan Drag & Drop, serta halaman Profil & Bantuan.\n\n"
                "Di sisi Admin, terdapat halaman Dashboard Analytics, Live Board Monitoring mahasiswa secara live, serta halaman System Audit Logging.\""
            ),
            "tip": "Tunjukkan kelengkapan modul aplikasi dari Landing Page hingga Admin Control Panel."
        },
        {
            "num": "SLIDE 6: 5. STRUKTUR DATABASE (SUNNAH / OPSIONAL)",
            "speaker": "MUHAMMAD ILHAM HABIBALLAH",
            "speaker_color": COLOR_ILHAM,
            "script": (
                "\"Selanjutnya untuk struktur basis data MySQL, meskipun opsional, TaskMate menggunakan 5 tabel relasional utama:\n"
                "1. Tabel mahasiswas: Menyimpan profil dan akun login mahasiswa.\n"
                "2. Tabel tasks: Menyimpan kartu tugas, status Kanban, prioritas, deadline, dan data checklist.\n"
                "3. Tabel activities: Mencatat audit log aktivitas transaksi pengerjaan tugas.\n"
                "4. Tabel task_reminders: Mencatat riwayat status pengiriman pengingat WhatsApp.\n"
                "5. Tabel admins: Menyimpan kredensial akun administrator pengelola.\n\n"
                "Relasi antar tabel terikat secara One-to-Many dengan fitur Cascade Delete untuk menjamin integritas data.\""
            ),
            "tip": "Jelaskan fungsi masing-masing tabel database secara singkat dan padat."
        },
        {
            "num": "SLIDE 7: 6. ALUR KERJA APLIKASI & LAYANAN PUSAT BANTUAN 🌟 [SLIDE FOKUS]",
            "speaker": "GIANJAR NUGRAHA",
            "speaker_color": COLOR_GIANJAR,
            "script": (
                "\"Bapak/Ibu Dewan Penguji, pada bagian ini kami menyajikan Alur Kerja Operasional Utama Sistem dan Layanan Pusat Bantuan:\n\n"
                "Untuk Alur Operasional, sistem berjalan dalam 4 tahap:\n"
                "• Pertama, Registrasi & Login Mahasiswa dengan NIM & Password (Bcrypt).\n"
                "• Kedua, Pembuatan Tugas Baru via AJAX POST yang masuk ke status 'TO DO'.\n"
                "• Ketiga, Siklus Kanban visual dengan Drag & Drop SortableJS, serta auto-done saat checklist 100% centang.\n"
                "• Keempat, Administrator memicu command PHP Artisan `tasks:send-deadline-reminders` untuk mengirim pesan pengingat deadline sisa H-5 & H-2 secara manual via Fonnte WA API.\n\n"
                "Selain itu, kami menyediakan **Layanan Bantuan Hubungi Admin** di mana mahasiswa dapat mengakses Pusat Bantuan dan terhubung langsung ke chat WhatsApp Admin ULBI untuk konsultasi kendala sistem secara real-time.\""
            ),
            "tip": "⭐ SLIDE PALING UTAMA! Runtunkan penjelasan alur operasional dan jelaskan kegunaan halaman Bantuan Hubungi Admin dengan jelas."
        },
        {
            "num": "SLIDE 8: 7. HASIL DARI APLIKASI & SIMPULAN",
            "speaker": "MUHAMMAD ILHAM HABIBALLAH",
            "speaker_color": COLOR_ILHAM,
            "script": (
                "\"Berdasarkan hasil analisis, perancangan, dan pengujian aplikasi TaskMate v2.0, dapat disimpulkan bahwa:\n"
                "1. Aplikasi telah berhasil diimplementasikan 100% berbasis web dengan Laravel, Tailwind CSS, dan MySQL.\n"
                "2. Papan kerja Kanban visual terbukti mempermudah mahasiswa mengelola dan memantau progres tugas secara terstruktur.\n"
                "3. Fitur pengiriman pengingat manual via WhatsApp Gateway terbukti membantu menekan keterlambatan pengumpulan tugas.\n"
                "4. Dasbor Admin memberikan transparansi monitoring keaktifan mahasiswa dan audit log secara real-time.\""
            ),
            "tip": "Tegaskan kesimpulan bahwa aplikasi telah 100% sukses dibangun dan siap dioperasikan."
        },
        {
            "num": "SLIDE 9: PENUTUP, Q&A & DEMO LIVE APLIKASI",
            "speaker": "ILHAM & GIANJAR (BERSAMA)",
            "speaker_color": COLOR_PRIMARY,
            "script": (
                "ILHAM:\n"
                "\"Demikian pemaparan laporan Proyek 1 dari Kelompok 26. Terima kasih banyak atas perhatian dan waktu Bapak/Ibu Dewan Penguji serta Dosen Pembimbing kami.\"\n\n"
                "GIANJAR:\n"
                "\"Selanjutnya, izinkan kami mendemonstrasikan aplikasi TaskMate secara langsung dan memohon arahan serta masukan dari Bapak/Ibu pada sesi tanya jawab. Terima kasih. Wassalamu’alaikum Warahmatullahi Wabarakatuh.\""
            ),
            "tip": "Langsung beralih ke browser web untuk mendemonstrasikan fitur-fitur aplikasi secara live."
        }
    ]

    for item in slides_detail:
        h_slide = doc.add_heading(level=2)
        r_s = h_slide.add_run(item["num"])
        r_s.font.name = "Arial"
        r_s.font.size = Pt(12.5)
        r_s.font.bold = True
        r_s.font.color.rgb = COLOR_DARK

        p_sp = doc.add_paragraph()
        r_sp = p_sp.add_run(f"🗣️ PEMBICARA: {item['speaker']}")
        r_sp.font.bold = True
        r_sp.font.size = Pt(11)
        r_sp.font.color.rgb = item["speaker_color"]

        p_sc = doc.add_paragraph()
        p_sc.paragraph_format.left_indent = Inches(0.2)
        r_sc = p_sc.add_run(item["script"])
        r_sc.font.size = Pt(11)

        p_tp = doc.add_paragraph()
        p_tp.paragraph_format.left_indent = Inches(0.2)
        r_tlab = p_tp.add_run("💡 Tips Penyampaian Sidang: ")
        r_tlab.font.bold = True
        r_tlab.font.size = Pt(9.5)
        r_tlab.font.color.rgb = COLOR_MUTED
        r_tp = p_tp.add_run(item["tip"])
        r_tp.font.italic = True
        r_tp.font.size = Pt(9.5)
        r_tp.font.color.rgb = COLOR_MUTED

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # GUIDANCE FOR LIVE DEMO
    doc.add_page_break()
    h_demo = doc.add_heading(level=1)
    r_hd = h_demo.add_run("💻 PANDUAN DEMO APLIKASI LANGSUNG SAAT SIDANG")
    r_hd.font.name = "Arial"
    r_hd.font.size = Pt(14)
    r_hd.font.bold = True
    r_hd.font.color.rgb = COLOR_PRIMARY

    demo_steps = [
        ("1. Persiapan Server", "Buka terminal, jalankan `php artisan serve` dan `npm run dev`. Buka browser di http://127.0.0.1:8000."),
        ("2. Demo Landing Page Publik", "Tunjukkan kartu statistik real-time di halaman utama (Total Mahasiswa, Total Tugas, Tugas Selesai)."),
        ("3. Demo Workspace Mahasiswa", "Login sebagai Mahasiswa -> Tunjukkan cara menambah tugas di kolom To Do -> Drag & Drop kartu ke Doing/Done -> Centang checklist 100% untuk demo fitur Auto-Done. Tunjukkan juga halaman Bantuan Hubungi Admin."),
        ("4. Demo Admin Control Panel", "Login sebagai Admin -> Tunjukkan tabel kelola mahasiswa -> Klik 'Lihat Board' untuk live monitoring Kanban mahasiswa -> Tunjukkan Audit Log Aktivitas.")
    ]

    for title, desc in demo_steps:
        p = doc.add_paragraph()
        r = p.add_run(f"• {title}: ")
        r.font.bold = True
        r.font.color.rgb = COLOR_DARK
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)

    output_path = "NASKAH_SIDANG_KELOMPOK_26.docx"
    doc.save(output_path)
    print(f"Word document saved to {os.path.abspath(output_path)}")

if __name__ == "__main__":
    create_naskah_word()
